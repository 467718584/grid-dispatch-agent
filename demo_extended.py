"""
扩展Skill演示 - 文件系统操作 & 代码执行
"""
import asyncio
import sys
import os
import subprocess
from datetime import datetime

sys.path.insert(0, '.')

from src.agent import GridAgent
from src.skill.base import BaseSkill


# ============ 文件系统Skill ============

class FileSystemSkill(BaseSkill):
    """文件系统操作Skill"""
    
    @property
    def name(self) -> str:
        return "file_system"
    
    @property
    def description(self) -> str:
        return "文件系统操作：读取、写入、列表、创建目录、删除、复制、移动"
    
    @property
    def input_schema(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "operation": {
                    "type": "string",
                    "enum": ["read", "write", "list", "mkdir", "delete", "copy", "move", "exists"],
                    "description": "操作类型"
                },
                "path": {"type": "string", "description": "文件/目录路径"},
                "content": {"type": "string", "description": "写入内容（write操作时）"},
                "dest": {"type": "string", "description": "目标路径（copy/move时）"},
                "filter": {"type": "string", "description": "文件过滤（list时，如*.py）"}
            },
            "required": ["operation", "path"]
        }
    
    async def execute(self, params, context):
        operation = params.get("operation")
        path = params.get("path")
        content = params.get("content", "")
        dest = params.get("dest")
        file_filter = params.get("filter")
        
        results = {}
        
        if operation == "exists":
            results["exists"] = os.path.exists(path)
            results["is_file"] = os.path.isfile(path) if results["exists"] else None
            results["is_dir"] = os.path.isdir(path) if results["exists"] else None
            
        elif operation == "read":
            if os.path.isfile(path):
                with open(path, 'r', encoding='utf-8') as f:
                    results["content"] = f.read()
                results["size"] = os.path.getsize(path)
            else:
                results["error"] = f"文件不存在: {path}"
                
        elif operation == "write":
            os.makedirs(os.path.dirname(path) if os.path.dirname(path) else '.', exist_ok=True)
            with open(path, 'w', encoding='utf-8') as f:
                f.write(content)
            results["success"] = True
            results["size"] = len(content)
            
        elif operation == "list":
            if os.path.isdir(path):
                items = os.listdir(path)
                if file_filter:
                    import fnmatch
                    items = [i for i in items if fnmatch.fnmatch(i, file_filter)]
                results["items"] = items
                results["count"] = len(items)
            else:
                results["error"] = f"目录不存在: {path}"
                
        elif operation == "mkdir":
            os.makedirs(path, exist_ok=True)
            results["success"] = True
            
        elif operation == "delete":
            if os.path.isfile(path):
                os.remove(path)
                results["success"] = True
            elif os.path.isdir(path):
                import shutil
                shutil.rmtree(path)
                results["success"] = True
            else:
                results["error"] = f"路径不存在: {path}"
                
        elif operation == "copy":
            import shutil
            shutil.copy2(path, dest)
            results["success"] = True
            results["dest"] = dest
            
        elif operation == "move":
            import shutil
            shutil.move(path, dest)
            results["success"] = True
            results["dest"] = dest
        
        return results


class FileOrganizeSkill(BaseSkill):
    """文件整理Skill - 按类型/日期整理文件"""
    
    @property
    def name(self) -> str:
        return "file_organize"
    
    @property
    def description(self) -> str:
        return "文件整理：按扩展名或日期分类整理文件"
    
    async def execute(self, params, context):
        import shutil
        import os
        
        source_dir = params.get("source_dir")
        target_dir = params.get("target_dir", source_dir + "_organized")
        organize_by = params.get("by", "ext")  # ext or date
        
        os.makedirs(target_dir, exist_ok=True)
        moved = []
        
        for filename in os.listdir(source_dir):
            src_path = os.path.join(source_dir, filename)
            if not os.path.isfile(src_path):
                continue
            
            if organize_by == "ext":
                ext = os.path.splitext(filename)[1][1:] or "no_ext"
                subdir = ext.upper()
            else:
                mtime = os.path.getmtime(src_path)
                date = datetime.fromtimestamp(mtime).strftime("%Y-%m")
                subdir = date
            
            dest_subdir = os.path.join(target_dir, subdir)
            os.makedirs(dest_subdir, exist_ok=True)
            dest_path = os.path.join(dest_subdir, filename)
            shutil.move(src_path, dest_path)
            moved.append({"file": filename, "to": subdir})
        
        return {"moved": moved, "total": len(moved), "target_dir": target_dir}


# ============ Word生成Skill ============

class WordGenerateSkill(BaseSkill):
    """Word文档生成Skill"""
    
    @property
    def name(self) -> str:
        return "word_generate"
    
    @property
    def description(self) -> str:
        return "生成Word文档（.docx），支持标题、段落、表格、列表"
    
    @property
    def input_schema(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "文档标题"},
                "body": {"type": "array", "description": "内容块列表"},
                "output_path": {"type": "string", "description": "输出文件路径"}
            }
        }
    
    async def execute(self, params, context):
        title = params.get("title", "无标题文档")
        body = params.get("body", [])
        output_path = params.get("output_path", "output.docx")
        
        # 检查python-docx是否可用
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return {
                "success": False,
                "error": "需要安装python-docx: pip install python-docx",
                "install_cmd": "pip install python-docx"
            }
        
        doc = Document()
        
        # 标题
        doc.add_heading(title, 0)
        
        # 内容块
        for block in body:
            block_type = block.get("type", "paragraph")
            
            if block_type == "heading":
                doc.add_heading(block.get("text", ""), block.get("level", 1))
            elif block_type == "paragraph":
                doc.add_paragraph(block.get("text", ""))
            elif block_type == "table":
                rows = block.get("rows", [])
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    for i, row_data in enumerate(rows):
                        row = table.rows[i]
                        for j, cell_data in enumerate(row_data):
                            row.cells[j].text = str(cell_data)
            elif block_type == "list":
                for item in block.get("items", []):
                    doc.add_paragraph(item, style='List Bullet')
        
        # 保存
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        doc.save(output_path)
        
        return {
            "success": True,
            "output_path": output_path,
            "title": title,
            "blocks_count": len(body)
        }


# ============ 代码执行Skill ============

class CodeExecSkill(BaseSkill):
    """代码执行Skill - 支持Python/JS/Shell"""
    
    @property
    def name(self) -> str:
        return "code_exec"
    
    @property
    def description(self) -> str:
        return "执行代码：Python/JS/Shell，返回执行结果"
    
    @property
    def input_schema(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "language": {
                    "type": "string",
                    "enum": ["python", "javascript", "node", "shell", "bash"],
                    "description": "代码语言"
                },
                "code": {"type": "string", "description": "要执行的代码"},
                "timeout": {"type": "integer", "description": "超时时间（秒）"}
            },
            "required": ["language", "code"]
        }
    
    async def execute(self, params, context):
        language = params.get("language", "python")
        code = params.get("code", "")
        timeout = params.get("timeout", 30)
        
        # 构建执行命令
        if language in ["python", "py"]:
            cmd = ["python3", "-c", code]
        elif language in ["javascript", "node"]:
            cmd = ["node", "-e", code]
        elif language in ["shell", "bash"]:
            cmd = ["bash", "-c", code]
        else:
            return {"error": f"不支持的语言: {language}"}
        
        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=timeout
            )
            
            return {
                "success": result.returncode == 0,
                "stdout": result.stdout,
                "stderr": result.stderr,
                "returncode": result.returncode,
                "language": language
            }
        except subprocess.TimeoutExpired:
            return {"error": f"执行超时（{timeout}秒）", "language": language}
        except Exception as e:
            return {"error": str(e), "language": language}


class CodeFileExecSkill(BaseSkill):
    """代码文件执行Skill - 执行.py/.js/.ts文件"""
    
    @property
    def name(self) -> str:
        return "code_file_exec"
    
    @property
    def description(self) -> str:
        return "执行代码文件：.py/.js/.ts文件，带参数支持"
    
    async def execute(self, params, context):
        file_path = params.get("file_path")
        args = params.get("args", [])
        timeout = params.get("timeout", 60)
        
        if not os.path.isfile(file_path):
            return {"error": f"文件不存在: {file_path}"}
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".py":
            cmd = ["python3", file_path] + args
        elif ext in [".js", ".mjs"]:
            cmd = ["node", file_path] + args
        elif ext == ".ts":
            # 需要ts-node
            cmd = ["npx", "ts-node", file_path] + args
        else:
            return {"error": f"不支持的文件类型: {ext}"}
        
        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=timeout,
                cwd=os.path.dirname(file_path) or "."
            )
            
            return {
                "success": result.returncode == 0,
                "stdout": result.stdout,
                "stderr": result.stderr,
                "returncode": result.returncode,
                "file": file_path
            }
        except subprocess.TimeoutExpired:
            return {"error": f"执行超时（{timeout}秒）", "file": file_path}
        except Exception as e:
            return {"error": str(e), "file": file_path}


# ============ 演示 ============

async def demo_file_system():
    """文件系统操作演示"""
    print("\n" + "="*50)
    print("演示1: 文件系统操作")
    print("="*50)
    
    agent = GridAgent(llm_url="http://localhost:8000/v1")
    agent.register_skill(FileSystemSkill())
    agent.set_flow(["file_system"])
    
    import tempfile
    import os
    
    # 创建测试目录和文件
    test_dir = tempfile.mkdtemp()
    
    # 测试写文件
    print("\n[Step 1] 写入文件...")
    result = await agent.execute(
        "写入测试数据",
        params={
            "operation": "write",
            "path": os.path.join(test_dir, "test.txt"),
            "content": "Hello Grid Agent!\nLine 2\nLine 3"
        }
    )
    print(f"  写文件: {result}")
    
    # 测试读文件
    print("\n[Step 2] 读取文件...")
    result = await agent.execute(
        "读取测试数据",
        params={
            "operation": "read",
            "path": os.path.join(test_dir, "test.txt")
        }
    )
    print(f"  读文件: {result}")
    
    # 测试列举文件
    print("\n[Step 3] 列举目录...")
    result = await agent.execute(
        "列举文件",
        params={
            "operation": "list",
            "path": test_dir
        }
    )
    print(f"  目录列表: {result}")
    
    # 清理
    import shutil
    shutil.rmtree(test_dir)
    print("\n  (测试目录已清理)")


async def demo_file_organize():
    """文件整理演示"""
    print("\n" + "="*50)
    print("演示2: 文件整理")
    print("="*50)
    
    agent = GridAgent(llm_url="http://localhost:8000/v1")
    agent.register_skill(FileSystemSkill())
    agent.register_skill(FileOrganizeSkill())
    
    import tempfile
    import os
    
    # 创建测试目录
    test_dir = tempfile.mkdtemp()
    target_dir = test_dir + "_organized"
    
    # 创建测试文件
    for ext in ["txt", "py", "js", "docx", "pdf"]:
        for i in range(2):
            filename = f"file_{i+1}.{ext}"
            with open(os.path.join(test_dir, filename), 'w') as f:
                f.write(f"Content of {filename}")
    
    print(f"\n[原始目录内容] {test_dir}")
    print(os.listdir(test_dir))
    
    # 整理文件
    print("\n[执行按扩展名整理...]")
    agent.set_flow(["file_organize"])  # FileOrganizeSkill需要自己的Flow
    result = await agent.execute(
        "整理文件",
        params={
            "source_dir": test_dir,
            "target_dir": target_dir,
            "by": "ext"
        }
    )
    print(f"  整理结果: {result}")
    
    # 查看整理后的结果
    print(f"\n[整理后目录结构] {target_dir}")
    for root, dirs, files in os.walk(target_dir):
        level = root.replace(target_dir, '').count(os.sep)
        indent = ' ' * 2 * level
        print(f"{indent}{os.path.basename(root)}/")
        subindent = ' ' * 2 * (level + 1)
        for file in files:
            print(f"{subindent}{file}")
    
    # 清理
    import shutil
    shutil.rmtree(test_dir)
    shutil.rmtree(target_dir)
    print("\n  (测试目录已清理)")


async def demo_word_generate():
    """Word生成演示"""
    print("\n" + "="*50)
    print("演示3: Word文档生成")
    print("="*50)
    
    agent = GridAgent(llm_url="http://localhost:8000/v1")
    agent.register_skill(WordGenerateSkill())
    
    import tempfile
    output_path = os.path.join(tempfile.gettempdir(), "grid_report.docx")
    
    result = await agent.execute(
        "生成报告",
        params={
            "title": "电网调度日报",
            "body": [
                {"type": "heading", "text": "一、概述", "level": 1},
                {"type": "paragraph", "text": "本日电网运行平稳，各项指标正常。"},
                {"type": "heading", "text": "二、负荷情况", "level": 1},
                {"type": "table", "rows": [["站点", "负荷(MW)"], ["A站", "120"], ["B站", "85"]]},
                {"type": "heading", "text": "三、建议", "level": 1},
                {"type": "list", "items": ["继续保持监控", "关注备用容量", "必要时调整调度"]}
            ],
            "output_path": output_path
        }
    )
    
    if result.get("status") == "success":
        data = result.get("data", {})
        inner_result = data.get("results", {}).get("word_generate", {})
        if inner_result.get("success"):
            print(f"  ✅ Word文档已生成: {inner_result.get('output_path')}")
        elif inner_result.get("error"):
            print(f"  ℹ️ {inner_result.get('error')}")
            print(f"  安装命令: {inner_result.get('install_cmd')}")
    else:
        print(f"  结果: {result}")


async def demo_code_exec():
    """代码执行演示"""
    print("\n" + "="*50)
    print("演示4: 代码执行")
    print("="*50)
    
    agent = GridAgent(llm_url="http://localhost:8000/v1")
    agent.register_skill(CodeExecSkill())
    agent.register_skill(CodeFileExecSkill())
    
    # 执行Python代码
    print("\n[Step 1] 执行Python代码...")
    result = await agent.execute(
        "执行Python计算",
        params={
            "language": "python",
            "code": "print('Hello from Python!'); print(2**10); import os; print(os.getcwd())"
        }
    )
    inner_result = result.get("data", {}).get("results", {}).get("code_exec", {})
    print(f"  stdout: {inner_result.get('stdout')}")
    print(f"  success: {inner_result.get('success')}")
    
    # 执行Shell命令
    print("\n[Step 2] 执行Shell命令...")
    result = await agent.execute(
        "执行Shell",
        params={
            "language": "shell",
            "code": "echo 'Current date:' && date && echo 'Files:' && ls -la ~/.openclaw/workspace/workspace-jisu/grid-dispatch-agent/ | head -5"
        }
    )
    inner_result = result.get("data", {}).get("results", {}).get("code_exec", {})
    print(f"  stdout: {inner_result.get('stdout')}")


async def demo_code_file_exec():
    """代码文件执行演示"""
    print("\n" + "="*50)
    print("演示5: 代码文件执行")
    print("="*50)
    
    agent = GridAgent(llm_url="http://localhost:8000/v1")
    agent.register_skill(CodeExecSkill())
    agent.register_skill(CodeFileExecSkill())
    
    import tempfile
    
    # 创建Python测试文件
    test_file = tempfile.mktemp(suffix=".py")
    with open(test_file, 'w') as f:
        f.write('''
import sys
print("Hello from file!")
print(f"Arguments: {sys.argv}")

# 简单计算
result = sum(range(1, 101))
print(f"Sum 1-100 = {result}")
''')
    
    print(f"\n[执行文件] {test_file}")
    result = await agent.execute(
        "执行Python文件",
        params={
            "file_path": test_file,
            "args": ["arg1", "arg2"]
        }
    )
    
    inner_result = result.get("data", {}).get("results", {}).get("code_file_exec", {})
    print(f"  stdout: {inner_result.get('stdout')}")
    print(f"  returncode: {inner_result.get('returncode')}")
    
    # 清理
    os.remove(test_file)


# ============ 运行所有演示 ============

if __name__ == "__main__":
    import os
    print("\n🟢 Grid Dispatch Agent - 扩展能力演示")
    print("="*50)
    
    asyncio.run(demo_file_system())
    asyncio.run(demo_file_organize())
    asyncio.run(demo_word_generate())
    asyncio.run(demo_code_exec())
    asyncio.run(demo_code_file_exec())
    
    print("\n" + "="*50)
    print("✅ 扩展能力演示完成")
    print("="*50)
    print("\n📋 总结：所有这些能力都通过Skill扩展实现，无需MCP！")
    print("   Skill就是Agent的'工具箱'，需要什么就注册什么。")